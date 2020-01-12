from tkinter import *
from tkinter import messagebox as ms

import sqlite3

from typing import List, Any

from PIL import ImageTk, Image

from docx import Document
from docx.shared import Inches

import os
# make database and users (if not exists already) table at programme start up
users = sqlite3.connect("users.db")
images = sqlite3.connect("images.db")
stories = sqlite3.connect("stories.db")



# users.cursor().execute("drop table users")
# images.cursor().execute("drop table images")
# stories.cursor().execute("drop table stories")


users.cursor().execute(
    "CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY ,password TEXT, age INT, role TEXT, child TEXT, class TEXT)"
)

images.cursor().execute(
    "CREATE TABLE IF NOT EXISTS images (username TEXT PRIMARY KEY ,images TEXT)"
)

stories.cursor().execute(
    "CREATE TABLE IF NOT EXISTS stories (username TEXT PRIMARY KEY, story TEXT)"
)

# conn.cursor().execute("UPDATE users SET story=NULL WHERE username='s'")

# conn.cursor().execute("UPDATE users SET class='yud' WHERE username='s'")

# conn.cursor().execute("UPDATE users SET images='img1.jpeg img2.jpeg img3.jpeg img4.jpeg' WHERE username='s'")

# print(conn.cursor().execute('SELECT child FROM users WHERE username = "p"').fetchone())


# print(conn.cursor().execute('SELECT * FROM users').fetchall())

print(
    users.cursor()
    .execute("SELECT username, role, child, class FROM users")
    .fetchall()
)

print(
    images.cursor()
    .execute("SELECT username ,images from images")
    .fetchall()
)

print(
    stories.cursor()
    .execute("SELECT * from stories")
    .fetchall()
)


# print(conn.cursor().execute('SELECT * FROM classes').fetchall())

TEACHER = "teacher"
PARENT = "parent"
STUDENT = "student"

imgs = {}




# Login Function
    def login(self):
        # Establish Connection
        # with sqlite3.connect('quit.db') as db:
        c = self.users_db.cursor()
        # Find user If there is any take proper action
        find_user = "SELECT * FROM users WHERE username = ? and password = ?"
        c.execute(find_user, [(self.username.get()), (self.password.get())])
        result = c.fetchone()
        if not result:
            ms.showerror("אוי", "שם המשתמש או הסיסמא לא תקינים")
            return
        self.logf.pack_forget()
        # self.head['text'] = f'Logged In As\n {self.username.get()}'
        # self.head['pady'] = 150
        username = result[0]
        role: str = result[3].lower()
        if role == TEACHER:
            c.execute(
                "SELECT class FROM users WHERE username=?", (self.username.get(),)
            )
            self.classname.set(c.fetchone()[0])
            self.show_teacher_frame()
        elif role == PARENT:
            self.setup_parent()
            self.show_parent_frame()
            # show parents frame
            pass
        elif role == STUDENT:
            c.execute(
                "SELECT class FROM users WHERE username=?", (self.username.get(),)
            )
            self.classname.set(c.fetchone()[0])
            self.show_student_frame()
        else:
            print(f"לא מכיר את התפקיד {username} - {role}")
            return

    def new_user(self):
        # Establish Connection
        users_cursor = self.users_db.cursor()
        imagess_cursor = self.images_db.cursor()
        stories_cursor = self.stories_db.cursor()

        # Find Existing username if any take proper action
        # find_user = "SELECT * FROM users WHERE username = ?"
        # users_cursor.execute(find_user, [(self.username.get())])
        # if users_cursor.fetchall():
        #     ms.showerror("Error!", "Username Taken Try a Diffrent One.")
        #     return

        if self.n_role.get() not in [TEACHER, STUDENT, PARENT]:
            ms.showerror(
                "תקלה!",
                f"התפקיד {self.n_role.get()} לא קיים",
            )
            return

        # Create New Account
        insert_users = "INSERT INTO users(username, password, age, role) VALUES(?,?,?,?)"
        insert_images = "INSERT INTO images(username) VALUES(?)"
        insert_stories = "INSERT INTO stories(username) VALUES(?)"
        try:
            users_cursor.execute(
            insert_users,
            [
                (self.n_username.get()),
                (self.n_password.get()),
                self.n_age.get(),
                self.n_role.get().lower(),
            ],
        )
        except Exception:
            return ms.showerror("!תקלה", "שם המשתמש קיים")
        imagess_cursor.execute(
            insert_images,
            [
                (self.n_username.get()),
            ],
        )
        stories_cursor.execute(
            insert_stories,
            [
                (self.n_username.get()),
            ],
        )
        self.users_db.commit()
        self.images_db.commit()
        self.stories_db.commit()
        ms.showinfo("בוצע", "המשתמש נוצר בהצלחה")
        self.login_frame()

    # Frame Packing Methords
    def login_frame(self):
        self.username.set("")
        self.password.set("")
        self.crf.pack_forget()
        self.teacher_frame.pack_forget()
        self.head["text"] = "התחברות"
        self.logf.pack()

    def create_acc_frame(self):
        self.n_username.set("")
        self.n_password.set("")
        self.logf.pack_forget()
        self.logf.pack_forget()
        self.head["text"] = "יצירת משתמש"
        self.crf.pack()


    def show_teacher_frame(self):
        self.parents_suggestions.pack_forget()
        self.logf.pack_forget()
        self.lesson.pack_forget()
        # self.head.pack_forget()
        self.head["text"] = f"ברוכים הבאים למסך המורה {self.username.get()}"
        self.teacher_frame.pack()


    def show_student_frame(self):
        if self.get_images_for_student() == -1:
            self.logf.pack()
            return
        self.logf.pack_forget()
        self.student_frame.pack()
        self.head["text"] = "הינה הקלפים לשיעורינו. "
        for i in range(4):
            Label(self.student_frame, image=imgs[self.images[i]]).grid(row=0, column=i)

    def show_parent_frame(self):
        self.suggest.pack_forget()
        self.head["text"] = "מסך הורה"
        self.parent_frame.pack()

    def suggest_to_teacher(self):
        if self.student.get() == "None":
            ms.showerror("תקלה", "לא הגדרת ילד")
            return
        c = self.users_db.cursor()
        child_class = "SELECT class FROM users WHERE username = ?"
        classname = c.execute(child_class, (self.student.get(),)).fetchone()
        self.head["text"] = "הצע תמונות למורה של ילדך"
        self.parent_frame.pack_forget()
        self.suggest.pack()

    def start_lesson(self):
        self.logf.pack_forget()
        self.teacher_frame.pack_forget()
        self.lesson.pack()
        self.head["text"] = "בחר קלפים לשיעור "
        img = None
        i = 0
        for child in self.lesson.children.values():
            if not isinstance(child, Button):
                continue
            i += 1
            if i > 12:
                return
            img = imgs[f"img{i}.jpeg"]
            child.configure(image=img)
            # child.image = img

    def image_button(self, button_id, frame):
        if len(self.images) >= 4:
            return ms.showerror("כבר נבחר", "כבר נבחרו 4 קלפים")
        button = list(frame.children.values())[button_id - 1]
        button.grid_remove()
        self.add_image(f"img{button_id}.jpeg")
        print(self.images)
        c = self.images_db.cursor()
        update_users = "INSERT OR REPLACE INTO images (images, username) VALUES (?, ?)"
        # update_classes = ('UPDATE classes SET images = ? WHERE name = ?')
        user_params = [(" ".join(self.images)), (self.username.get())]
        # class_params = [(' '.join(self.images)), (self.classname.get())]
        c.execute(update_users, user_params)
        # c.execute(update_classes, class_params)
        self.images_db.commit()

    def add_image(self, name: str):
        if name not in self.images:
            self.images.append(name)


    def create_class(self):
        c = self.users_db.cursor()
        c.execute("SELECT class FROM users WHERE username=?", (self.username.get(),))
        classname = c.fetchone()[0]
        print(classname)
        if classname:
            self.classname.set(classname)
            ms.showerror("טעות", "כבר יש לך כיתה!")
            return
        uname = self.username.get()
        classname = self.classname.get()
        if classname == "None":
            ms.showerror("טעות", "את חייבת לבחור שם לכיתה")
            return

        update_users = "UPDATE users SET class = ? WHERE username = ?"
        # insert_classes = ('INSERT INTO classes(name, teacher) VALUES (?, ?)')
        c.execute(update_users, [(classname), (uname)])
        # c.execute(insert_classes, [(classname), (uname)])
        self.users_db.commit()
        ms.showinfo("הצלחה", "הכיתה נוצרה")

    def add_student(self):
        if not self.classname.get():
            ms.showerror("טעות", "עדיין לא יצרת כיתה")
            return
        c = self.users_db.cursor()
        find_user = "SELECT * FROM users WHERE username = ? AND role = ?"
        c.execute(find_user, [(self.student.get()), (STUDENT)])
        student = c.fetchone()
        if not student:
            ms.showerror("Error!", "Student in that name is not found")
            return
        # Create New Account
        update = "UPDATE users SET class = ? WHERE username = ?"
        c.execute(update, [(self.classname.get()), (self.student.get())])
        self.users_db.commit()
        ms.showinfo(
            "הצלחה",
            f"התווסף {self.student.get()} לכיתה-  {self.classname.get()}",)


    def send_story(self):
        if not self.classname.get():
            ms.showerror("תקלה", "אתה לא בכיתה עדיין")
            return
        c = self.stories_db.cursor()
        find_story = "SELECT story FROM stories WHERE username = ?"
        c.execute(find_story, [(self.username.get()),])
        result, = c.fetchone()
        stories = str(result) + self.story.get() + "|"
        if result is None:
            stories = self.story.get() + "|"
        update = "UPDATE stories SET story = ? WHERE username = ?"
        c.execute(update, [(stories), (self.username.get())])
        self.stories_db.commit()
        print((c.execute("SELECT username, story FROM stories").fetchall()))
        ms.showinfo("הצלחה", f"הסיפור נשמר")
        self.story.set("")  # clear the story
     
    def get_images_for_student(self):
        users_cursor = self.users_db.cursor()
        images_cursor = self.images_db.cursor()
        find_user = "SELECT username FROM users WHERE class = ? AND role = ?"
        usr_tuple = users_cursor.execute(
            find_user, [(self.classname.get()), ("teacher")]
        ).fetchone()
        if not usr_tuple:
            ms.showerror("טעות", "המורה שלך לא יצר כיתה עדיין!")
            return -1
        find_images = f"SELECT images FROM images WHERE username = '{usr_tuple[0]}'"
        print(find_images)
        images, = images_cursor.execute(find_images).fetchone()
        if not images:
            ms.showerror("טעות", "המורה לא הגדיר שיעור עדיין")
            return -1

        print(len(images))
        
        for img in images.split(" "):
            self.add_image(img)
        print(self.images)
     
    def set_child(self):
        c = self.users_db.cursor()
        find_user = "SELECT * FROM users WHERE username = ?"
        student = self.student.get()
        res = c.execute(find_user, [(student)]).fetchone()
        print(res)
        if not res:
            return ms.showerror("טעות!", "ילד בשם זה לא קיים במערכת")
        update_child = "UPDATE users SET child = ? WHERE username = ?"
        c.execute(update_child, [(student), (self.username.get())])

        child_class = "SELECT class FROM users WHERE username = ?"
        classname = c.execute(child_class, (student,)).fetchone()
        self.classname.set(classname)
        update = "UPDATE users SET class = ? WHERE username = ?"
        c.execute(update, [(classname[0]), (self.username.get())])
        ms.showinfo("הצלחה!", "ילדך נשמר במערכת")

        self.users_db.commit()

    # setup parent child and class before showing interface
    def setup_parent(self):
        c = self.users_db.cursor()
        username = self.username.get()
        find_child = "SELECT child FROM users WHERE username = ?"
        child_class = "SELECT class FROM users WHERE username = ?"
        child = c.execute(find_child, (username,)).fetchone()
        classname = c.execute(child_class, (child[0],)).fetchone()
        print(child, classname)
        self.student.set(child[0])
        if classname:
            self.classname.set(classname[0])



    def view_parent_seuggestions(self):
        self.head["text"] = "ראה את הצעת ההורים לשיעור הבא שלך"
        self.teacher_frame.pack_forget()
        self.parents_suggestions.pack()
        c = self.users_db.cursor()
        find_parents = (
            "SELECT username FROM users WHERE class=? AND role=?"
        )
        parents = c.execute(
            find_parents, [(self.classname.get()), ("parent")]
        ).fetchall()
        print(parents)
        for parent in parents:
            if parent[0] not in self.parents:
                self.parents.append(parent[0])
            index = parents.index(parent)
            Button(
                self.parents_suggestions,
                text=parent[0],
                height=10,
                width=20,
                command=lambda i=index: self.show_suggestion(i),
            ).grid(row=index, column=index // 4)

    def show_suggestion(self, index):
        self.head["text"] = f" ההצעה של - {self.parents[index]} "
        self.suggest.pack_forget()
        c = self.images_db.cursor()
        parent = self.parents[index]
        find_suggestions = "SELECT images FROM images WHERE username=?"
        suggestions = c.execute(find_suggestions, (parent,)).fetchone()
        if not suggestions:
            ms.showerror(
                "תקלה", "ההורה הזה לא הגדיר ילד,או שלילד אין כיתה"
            )
            return
        # suggestions = list(suggestions)
        images = "".join(suggestions).split(" ")
        print(str(index) + " : " + str(images))
        # print(images)
        for img in images:
            # print(img)
            Label(self.parents_suggestions, image=imgs[str(img)]).grid(
                row=0, column=images.index(img)
            )

    def student_report(self):
        users_db = self.users_db.cursor()
        stories_db = self.stories_db.cursor()

        # we can use self.student.get(), but we prefer being accurate
        (child,) = users_db.execute(
            "SELECT child FROM users WHERE username=?", (self.username.get(),)
        ).fetchone()
        (stories_raw,) = stories_db.execute(
            "SELECT story FROM stories WHERE username=?", (child,)
        ).fetchone()
        if not stories_raw:
            ms.showerror("טעות", "לילדך אין סיפורים עדיין")
            return

        stories = stories_raw.split("|")  # stories are split with a '|'
        try:
            os.remove(f"{child}_report.docx")
        except Exception:
            pass
        report = Document()
        report.add_heading(f"{child}הסיפור של\n", 0)
        for story in stories:
            p = report.add_paragraph(f"סיפור מספר {stories.index(story) + 1}\n\n")
            p.add_run(story + "\n\n\n").italic = True

        report.save(f"{child}_report.docx")
        os.system(f"start {child}_report.docx")
    
    
    # Draw Widgets
    def widgets(self):
        self.head = Label(self.master, text="התחברות", font=("", 35), pady=10)
        self.head.pack()
        self.logf = Frame(self.master, padx=10, pady=10)
        Label(self.logf, text="שם משתמש: ", font=("", 20), pady=5, padx=5).grid(
            sticky=W
        )
        Entry(self.logf, textvariable=self.username, bd=5, font=("", 15)).grid(
            row=0, column=1
        )
        Label(self.logf, text="סיסמא: ", font=("", 20), pady=5, padx=5).grid(
            sticky=W
        )
        Entry(
            self.logf, textvariable=self.password, bd=5, font=("", 15), show="*"
        ).grid(row=1, column=1)
        Button(
            self.logf,
            text=" התחברות ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.login,
        ).grid()
        Button(
            self.logf,
            text=" יצירת משתמש ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.create_acc_frame,
        ).grid(row=2, column=1)
        self.logf.pack()

        self.crf = Frame(self.master, padx=10, pady=10)
        Label(self.crf, text="שם משתמש: ", font=("", 20), pady=10, padx=10).grid(
            sticky=W
        )
        Entry(self.crf, textvariable=self.n_username, bd=5, font=("", 15)).grid(
            row=0, column=1
        )
        Label(self.crf, text="סיסמא: ", font=("", 20), pady=5, padx=5).grid(sticky=W)
        Entry(
            self.crf, textvariable=self.n_password, bd=5, font=("", 15), show="*"
        ).grid(row=1, column=1)
        Label(self.crf, text="תפקיד: ", font=("", 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.crf, textvariable=self.n_role, bd=5, font=("", 15)).grid(
            row=2, column=1
        )
        Label(self.crf, text="גיל: ", font=("", 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.crf, textvariable=self.n_age, bd=5, font=("", 15)).grid(
            row=3, column=1
        )
        Button(
            self.crf,
            text=" יצירת משתמש ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.new_user,
        ).grid()
        Button(
            self.crf,
            text="חזור להתחברות ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.login_frame,
        ).grid(row=4, column=1)

        self.teacher_frame = Frame(self.master, padx=100, pady=100)
        Label(
            self.teacher_frame,
            text="שם התלמיד",
            font=("", 20),
            pady=5,
            padx=5,
        ).grid(sticky=W)
        Entry(self.teacher_frame, textvariable=self.student, bd=5, font=("", 15)).grid(
            row=0, column=1
        )
        Button(
            self.teacher_frame,
            text="צרף תלמיד ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.add_student,
        ).grid(row=0, column=2)

        Label(
            self.teacher_frame,
            text=" שם הכיתה - ",
            font=("", 20),
            pady=5,
            padx=5,
        ).grid(sticky=W)
        Entry(
            self.teacher_frame, textvariable=self.classname, bd=5, font=("", 15)
        ).grid(row=1, column=1)
        Button(
            self.teacher_frame,
            text=" צור כיתה ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.create_class,
        ).grid(row=1, column=2)

        Button(
            self.teacher_frame,
            text=" צרף תלמיד ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.add_student,
        ).grid(row=0, column=2)
        Button(
            self.teacher_frame,
            text="התחל שיעור",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.start_lesson,
        ).grid(row=2, column=0)
        Button(
            self.teacher_frame,
            text="ראה את הצעת ההורים",
            bd=3,
            font=("", 15),
            padx=0,
            pady=0,
            command=self.view_parent_seuggestions,
        ).grid(row=3, column=0)
        Button(
            self.teacher_frame,
            text="חזור להתחברות",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.login_frame,
        ).grid(row=10, column=0)

        self.lesson = Frame(self.master, padx=10, pady=10)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(1, self.lesson),
        ).grid(row=0, column=1)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(2, self.lesson),
        ).grid(row=0, column=2)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(3, self.lesson),
        ).grid(row=0, column=3)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(4, self.lesson),
        ).grid(row=0, column=4)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(5, self.lesson),
        ).grid(row=1, column=1)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(6, self.lesson),
        ).grid(row=1, column=2)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(7, self.lesson),
        ).grid(row=1, column=3)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(8, self.lesson),
        ).grid(row=1, column=4)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(9, self.lesson),
        ).grid(row=2, column=1)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(10, self.lesson),
        ).grid(row=2, column=2)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(11, self.lesson),
        ).grid(row=2, column=3)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            command=lambda: self.image_button(12, self.lesson),
        ).grid(row=2, column=4)
        Button(
            self.lesson,
            padx=5,
            pady=5,
            font=("", 20),
            command=self.show_teacher_frame,
            text="חזור למסך המורה הראשי ",
        ).grid(row=0, column=0)

        self.student_frame = Frame(self.master, padx=100, pady=100)
        Label(
            self.student_frame, text=" התחל כתיבת סיפור ", font=("", 20), pady=5, padx=5
        ).grid(row=1, sticky=W)
        Entry(self.student_frame, textvariable=self.story, bd=5, font=("", 15)).grid(
            row=1, column=1
        )
        # Text(self.student_frame, height=20, width=60).grid()
        Button(
            self.student_frame,
            text=" הוסף סיפור ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.send_story,
        ).grid(row=1, column=2)

        self.parent_frame = Frame(self.master, padx=100, pady=100)
        Label(
            self.parent_frame, text="שם הילד -  ", font=("", 20), pady=5, padx=5
        ).grid(row=1, sticky=W)
        Entry(self.parent_frame, textvariable=self.student, bd=5, font=("", 15)).grid(
            row=1, column=1
        )
        # Text(self.student_frame, height=20, width=60).grid()
        Button(
            self.parent_frame,
            text="הגדר ילד ",
            bd=3,
            font=("", 15),
            padx=5,
            pady=5,
            command=self.set_child,
        ).grid(row=1, column=2)

        Button(
            self.parent_frame,
            text="הצע למורה ",
            pady=5,
            padx=5,
            font=("", 15),
            height=5,
            width=15,
            command=self.suggest_to_teacher,
        ).grid(row=5, column=1)
        Button(
            self.parent_frame,
            text="הפק דוח תלמיד",
            pady=5,
            padx=5,
            font=("", 15),
            height=5,
            width=20,
            command=self.student_report,
        ).grid(row=6, column=1)

        self.suggest = Frame(self.master, padx=100, pady=100)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(1, self.suggest),
            image=imgs["img1.jpeg"],
        ).grid(row=0, column=1)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(2, self.suggest),
            image=imgs["img2.jpeg"],
        ).grid(row=0, column=2)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(3, self.suggest),
            image=imgs["img3.jpeg"],
        ).grid(row=0, column=3)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(4, self.suggest),
            image=imgs["img4.jpeg"],
        ).grid(row=0, column=4)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(5, self.suggest),
            image=imgs["img5.jpeg"],
        ).grid(row=1, column=1)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(6, self.suggest),
            image=imgs["img6.jpeg"],
        ).grid(row=1, column=2)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(7, self.suggest),
            image=imgs["img7.jpeg"],
        ).grid(row=1, column=3)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(8, self.suggest),
            image=imgs["img8.jpeg"],
        ).grid(row=1, column=4)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(9, self.suggest),
            image=imgs["img9.jpeg"],
        ).grid(row=2, column=1)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(10, self.suggest),
            image=imgs["img10.jpeg"],
        ).grid(row=2, column=2)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(11, self.suggest),
            image=imgs["img11.jpeg"],
        ).grid(row=2, column=3)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            command=lambda: self.image_button(12, self.suggest),
            image=imgs["img12.jpeg"],
        ).grid(row=2, column=4)
        Button(
            self.suggest,
            padx=5,
            pady=5,
            font=("", 20),
            command=self.show_parent_frame,
            text="חזור למסך המורה ",
        ).grid(row=0, column=0)

        self.parents_suggestions = Frame(self.master, padx=100, pady=100)
        Button(
            self.parents_suggestions,
            padx=5,
            pady=5,
            font=("", 20),
            command=self.show_teacher_frame,
            text="חזור למסך המורה ",
        ).grid(row=10, column=1)
